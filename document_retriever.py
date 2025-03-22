import os
import fitz
import docx

class DocumentRetriever:
    def __init__(self, directory):
        self.directory = directory
        self.documents = self.load_documents()

    def load_documents(self):
        documents = []
        for filename in os.listdir(self.directory):
            if filename.endswith(".txt"):
                with open(os.path.join(self.directory, filename), 'r') as file:
                    documents.append(file.read())
            elif filename.endswith(".pdf"):
                documents.append(self.load_pdf(os.path.join(self.directory, filename)))
            elif filename.endswith(".docx"):
                documents.append(self.load_docx(os.path.join(self.directory, filename)))
        return documents

    def load_pdf(self, filepath):
        doc = fitz.Document(filepath)
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def load_docx(self, filepath):
        doc = docx.Document(filepath)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    def retrieve(self):
        # For simplicity, return all documents
        return "\n\n".join(self.documents)